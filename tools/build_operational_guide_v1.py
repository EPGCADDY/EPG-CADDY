#!/usr/bin/env python3
"""Build the branded Golf Score Card GT user-operating guide.

Design basis:
- documents skill preset: compact_reference_guide
- header pattern: editorial_cover
- named brand override: Arial + Golf Score Card GT black/neon-green palette
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "GUIA_OPERATIVA_GOLF_SCORE_CARD_GT_V1.md"
OUTPUT = ROOT / "docs" / "GUIA_OPERATIVA_GOLF_SCORE_CARD_GT_V1.docx"
LOGO = ROOT / "assets" / "official-logos" / "golf-score-card-gt-official-master-1254.jpeg"

# compact_reference_guide tokens + named Golf Score Card GT brand override.
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1.0)
HEADER_DISTANCE = Inches(0.492)
FOOTER_DISTANCE = Inches(0.492)
CONTENT_DXA = 9360
LIST_MARKER_DXA = 270  # 0.1875 in
LIST_TEXT_DXA = 540  # 0.375 in
LIST_HANG_DXA = 270
BODY_FONT = "Arial"
BODY_SIZE = Pt(10.5)
BODY_AFTER = Pt(4)
BODY_LINE = 1.15
H1_BEFORE = Pt(0)
H1_AFTER = Pt(10)
H2_BEFORE = Pt(11)
H2_AFTER = Pt(5)
H3_BEFORE = Pt(8)
H3_AFTER = Pt(4)

BLACK = "050505"
NEON = "31FF00"
DARK_GREEN = "146600"
DEEP_GREEN = "0A3B00"
MUTED = "626A70"
LIGHT_GREEN = "F0FFE9"
LIGHT_GRAY = "F3F5F6"
SILVER = "C7CDD1"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cellless_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side: str, color: str, size: int = 14, space: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=BODY_FONT) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = size if hasattr(size, "pt") else Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color) if isinstance(color, str) else color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: Pt, color: str, bold: bool) -> None:
    style.font.name = BODY_FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style.font.size = size
    style.font.color.rgb = rgb(color)
    style.font.bold = bold


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, BODY_SIZE, BLACK, False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = BODY_AFTER
    normal.paragraph_format.line_spacing = BODY_LINE

    h1 = styles["Heading 1"]
    set_style_font(h1, Pt(16), NEON, True)
    h1.paragraph_format.space_before = H1_BEFORE
    h1.paragraph_format.space_after = H1_AFTER
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False

    h2 = styles["Heading 2"]
    set_style_font(h2, Pt(12.5), DARK_GREEN, True)
    h2.paragraph_format.space_before = H2_BEFORE
    h2.paragraph_format.space_after = H2_AFTER
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, Pt(11.5), DEEP_GREEN, True)
    h3.paragraph_format.space_before = H3_BEFORE
    h3.paragraph_format.space_after = H3_AFTER
    h3.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        set_style_font(style, BODY_SIZE, BLACK, False)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = BODY_LINE
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    if "Guide Callout" not in styles:
        callout = styles.add_style("Guide Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Guide Callout"]
    set_style_font(callout, Pt(10.5), BLACK, False)
    callout.paragraph_format.left_indent = Inches(0.16)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(7)
    callout.paragraph_format.line_spacing = 1.15

    if "Guide Kicker" not in styles:
        kicker = styles.add_style("Guide Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = styles["Guide Kicker"]
    set_style_font(kicker, Pt(9.5), DARK_GREEN, True)
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    kicker.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])
    set_run_font(run, size=Pt(8.5), color=MUTED)


def configure_section(section) -> None:
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("GOLF SCORE CARD GT  ·  GUÍA OPERATIVA")
    set_run_font(run, size=Pt(8.5), color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    label = fp.add_run("APP V306  ·  PÁGINA ")
    set_run_font(label, size=Pt(8.5), color=MUTED, bold=True)
    add_page_field(fp)


def add_image_alt(paragraph, description: str) -> None:
    doc_pr = paragraph._p.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", description)
        doc_pr[0].set("title", description)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(4.55))
    add_image_alt(p, "Logo oficial Golf Score Card GT")

    band = doc.add_paragraph()
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    band.paragraph_format.space_before = Pt(4)
    band.paragraph_format.space_after = Pt(12)
    band.paragraph_format.left_indent = Inches(0.40)
    band.paragraph_format.right_indent = Inches(0.40)
    set_cellless_shading(band, BLACK)
    run = band.add_run("GUÍA OPERATIVA")
    set_run_font(run, size=Pt(24), color=NEON, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("GOLF SCORE CARD GT")
    set_run_font(run, size=Pt(25), color=BLACK, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(22)
    run = sub.add_run("Juega · Dicta · Comprueba · La tarjeta hace el resto")
    set_run_font(run, size=Pt(12.5), color=DARK_GREEN, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run("EDICIÓN 1.0  ·  APP V306  ·  USUARIO FINAL")
    set_run_font(run, size=Pt(10), color=MUTED, bold=True)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.paragraph_format.space_after = Pt(20)
    run = meta2.add_run("Uso principal: iPhone  ·  De 1 a 6 jugadores")
    set_run_font(run, size=Pt(9.5), color=MUTED)

    purpose = doc.add_paragraph(style="Guide Callout")
    purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cellless_shading(purpose, LIGHT_GREEN)
    set_paragraph_border(purpose, "left", NEON, size=22, space=8)
    run = purpose.add_run(
        "Desde abrir la aplicación hasta cerrar, guardar y compartir una ronda, "
        "sin conocimientos técnicos."
    )
    set_run_font(run, size=Pt(11), color=BLACK, bold=True)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(18)
    run = date.add_run("24 DE AGOSTO DE 2026")
    set_run_font(run, size=Pt(9), color=MUTED, bold=True)
    doc.add_page_break()


def add_quick_index(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("MAPA DE LA GUÍA")
    set_run_font(run, size=Pt(9.5), color=DARK_GREEN, bold=True)

    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(12)
    run = h.add_run("Encuentra lo que necesitas en segundos")
    set_run_font(run, size=Pt(20), color=BLACK, bold=True)

    entries = [
        ("PARTE I", "Instalar, elegir modalidad y comenzar."),
        ("PARTE II", "Registrar jugadores en General y Stableford."),
        ("PARTE III", "Anotar, consultar y corregir durante la ronda."),
        ("PARTE IV", "Finalizar, guardar, compartir e Historial."),
        ("PARTE V", "Solución de atrancones y recuperación segura."),
        ("EMERGENCIA", "Una página para resolver lo esencial en el campo."),
        ("GLOSARIO", "Significado directo de cada concepto."),
    ]
    for label, detail in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.left_indent = Inches(0.12)
        set_paragraph_border(p, "left", NEON, size=18, space=6)
        r1 = p.add_run(f"{label}  ")
        set_run_font(r1, size=Pt(10.5), color=DARK_GREEN, bold=True)
        r2 = p.add_run(detail)
        set_run_font(r2, size=Pt(10.5), color=BLACK)
    doc.add_page_break()


def next_numbering_id(numbering) -> tuple[int, int]:
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    return (max(abstract_ids, default=-1) + 1, max(num_ids, default=0) + 1)


def create_real_list(doc: Document, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = next_numbering_id(numbering)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    lvl.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(LIST_TEXT_DXA))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(LIST_TEXT_DXA))
    ind.set(qn("w:hanging"), str(LIST_HANG_DXA))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "50")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if not ordered:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), BODY_FONT)
        fonts.set(qn("w:hAnsi"), BODY_FONT)
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.insert(0, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str, *, default_bold: bool = False, color: str = BLACK) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=BODY_SIZE, color=color, bold=default_bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=BODY_SIZE, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=Pt(10.5), color=DEEP_GREEN, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=BODY_SIZE, color=color, bold=default_bold)


def add_part_heading(doc: Document, text: str) -> None:
    # Use an explicit page break instead of `page_break_before` on Heading 1.
    # This renders consistently in Word and LibreOffice and keeps the branded
    # band inside the document margins on every section-opening page.
    # Parts II and IV deliberately continue on the current page: the preceding
    # topic is short and this avoids two nearly empty pages in a field guide.
    continues_current_page = text.startswith("PARTE II") or text.startswith("PARTE IV")
    if not continues_current_page:
        doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(14) if continues_current_page else Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.10)
    set_cellless_shading(p, BLACK)
    for side in ("top", "bottom", "left", "right"):
        set_paragraph_border(p, side, BLACK, size=20, space=0)
    run = p.add_run(text)
    set_run_font(run, size=Pt(16), color=NEON, bold=True)


def add_blockquote(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph(style="Guide Callout")
    set_cellless_shading(p, LIGHT_GREEN)
    set_paragraph_border(p, "left", NEON, size=22, space=7)
    joined = "\n".join(line.strip() for line in lines)
    add_inline(p, joined)


def body_lines(source_text: str) -> list[str]:
    lines = source_text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## CÓMO USAR ESTA GUÍA")
    return lines[start:]


def add_body(doc: Document, lines: list[str]) -> None:
    i = 0
    active_kind = None
    active_num_id = None
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped or stripped == "---":
            active_kind = None
            active_num_id = None
            i += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_blockquote(doc, quote_lines)
            active_kind = None
            active_num_id = None
            continue

        if stripped.startswith("# "):
            add_part_heading(doc, stripped[2:].strip())
            active_kind = None
            active_num_id = None
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[3:].strip(), default_bold=True, color=DARK_GREEN)
            active_kind = None
            active_num_id = None
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[4:].strip(), default_bold=True, color=DEEP_GREEN)
            active_kind = None
            active_num_id = None
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet_match or number_match:
            kind = "number" if number_match else "bullet"
            if active_kind != kind:
                active_num_id = create_real_list(doc, ordered=(kind == "number"))
                active_kind = kind
            p = doc.add_paragraph(style="List Number" if kind == "number" else "List Bullet")
            apply_num(p, active_num_id)
            add_inline(p, (number_match or bullet_match).group(1))
            i += 1
            continue

        active_kind = None
        active_num_id = None
        paragraph_parts = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate or candidate == "---" or candidate.startswith("#") or candidate.startswith(">"):
                break
            if re.match(r"^-\s+", candidate) or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_parts.append(candidate)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_parts))


def audit_document(doc: Document, source_text: str) -> None:
    required = [
        "1-# JUGADOR",
        "2-NOMBRE",
        "HASTA 6 JUGADORES",
        "3-OK",
        "HISTORIAL",
        "ATRÁS",
        "REGÍSTRATE",
        "SOLUCIÓN DE ATRANCONES",
        "GUÍA DE EMERGENCIA DE UNA PÁGINA",
        "Campos disponibles en todas las modalidades",
        "IN, OUT y TOTAL",
    ]
    missing = [item for item in required if item not in source_text]
    if missing:
        raise SystemExit(f"Missing required guide content: {missing}")
    retired_term = "BIBLIO" + "TECA"
    if re.search(rf"\b{retired_term}(?: DE TARJETAS)?\b", source_text, flags=re.I):
        raise SystemExit("Retired visible term detected in guide source")
    if re.search(r"\b(?:Jaime|Roberto|Ana|Eduardo)\b", source_text, flags=re.I):
        raise SystemExit("Real-person example detected in guide source")
    if "OUT, IN" in source_text:
        raise SystemExit("Incorrect OUT/IN order detected in guide source")
    fields = ["El Pulté", "Country Club", "San Isidro", "Mayan Golf", "Hacienda Nueva", "Alta Vista", "La Reunión"]
    if any(field not in source_text for field in fields):
        raise SystemExit("Universal course catalog is incomplete in guide source")
    if not LOGO.is_file():
        raise SystemExit("Official logo is missing")
    if len(doc.paragraphs) < 220:
        raise SystemExit("Guide appears incomplete")


def main() -> None:
    source_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    doc.core_properties.title = "Guía Operativa Golf Score Card GT"
    doc.core_properties.subject = "Manual operativo para usuario final"
    doc.core_properties.author = "Golf Score Card GT"
    doc.core_properties.keywords = "golf, score card, guía operativa, iPhone, Stableford"
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_cover(doc)
    add_quick_index(doc)
    add_body(doc, body_lines(source_text))
    audit_document(doc, source_text)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"PASS guide build: {OUTPUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
